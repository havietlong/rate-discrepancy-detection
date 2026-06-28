"""
Invoice Editor Module - True Word-like Document Editor
Converts PDF to editable document format and allows editing
"""

import streamlit as st
import fitz  # PyMuPDF
import pandas as pd
import re
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os

# Vietnam tax rates
SERVICE_CHARGE_RATE = 0.05
VAT_RATE = 0.08
TAX_RATE = 1.134

def pdf_to_docx(pdf_bytes):
    """
    Convert PDF to editable DOCX document
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    
    # Create a new Word document
    word_doc = Document()
    
    # Set page margins
    for section in word_doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        
        # Get text blocks with formatting
        blocks = page.get_text("dict")
        
        for block in blocks.get("blocks", []):
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if text:
                        # Create paragraph with formatting
                        p = word_doc.add_paragraph()
                        run = p.add_run(text)
                        
                        # Try to preserve formatting
                        if span.get("size", 0) > 12:
                            run.font.size = Pt(span.get("size", 11))
                        
                        if span.get("flags", 0) & 2**0:  # Bold
                            run.bold = True
                        
                        if span.get("flags", 0) & 2**1:  # Italic
                            run.italic = True
                        
                        # Check if it's a header or title
                        if "INFORMATION INVOICE" in text.upper():
                            run.font.size = Pt(16)
                            run.bold = True
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif "INVOICE" in text.upper():
                            run.font.size = Pt(14)
                            run.bold = True
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page break between pages (except last)
        if page_num < len(doc) - 1:
            word_doc.add_page_break()
    
    doc.close()
    return word_doc

def docx_to_pdf(word_doc):
    """
    Convert Word document back to PDF bytes
    """
    # Save docx to temp file
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
        word_doc.save(tmp_docx.name)
        docx_path = tmp_docx.name
    
    # Convert docx to pdf using pdf_converter (we'll use the built-in approach)
    try:
        from docx2pdf import convert
        pdf_path = docx_path.replace('.docx', '.pdf')
        convert(docx_path, pdf_path)
        
        with open(pdf_path, 'rb') as f:
            pdf_bytes = f.read()
        
        # Clean up temp files
        os.unlink(docx_path)
        os.unlink(pdf_path)
        
        return pdf_bytes
    except:
        # Fallback: return the docx as is
        with open(docx_path, 'rb') as f:
            docx_bytes = f.read()
        os.unlink(docx_path)
        return docx_bytes

def display_invoice_editor(pdf_bytes):
    """
    True Word-like document editor
    """
    st.subheader("📄 Invoice Editor")
    
    # Initialize session state for document
    if 'word_doc' not in st.session_state:
        st.session_state.word_doc = None
    if 'show_editor' not in st.session_state:
        st.session_state.show_editor = False
    
    # ===== TOOLBAR =====
    st.markdown("""
    <style>
    .word-toolbar {
        background: #f3f5f7;
        padding: 10px 15px;
        border-radius: 8px 8px 0 0;
        border: 1px solid #ddd;
        border-bottom: none;
        display: flex;
        flex-wrap: wrap;
        gap: 10px;
    }
    .word-toolbar button {
        background: white;
        border: 1px solid #ccc;
        border-radius: 4px;
        padding: 4px 12px;
        cursor: pointer;
        font-size: 13px;
    }
    .word-toolbar button:hover {
        background: #e8edf2;
    }
    .word-editor {
        border: 1px solid #ddd;
        border-radius: 0 0 8px 8px;
        padding: 20px;
        background: white;
        min-height: 500px;
    }
    .word-editor textarea {
        width: 100%;
        min-height: 500px;
        border: none;
        outline: none;
        font-family: 'Segoe UI', Arial, sans-serif;
        font-size: 14px;
        line-height: 1.6;
        padding: 10px;
        resize: vertical;
    }
    .word-editor textarea:focus {
        border: none;
        outline: none;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # ===== CONVERT PDF TO DOCUMENT =====
    if not st.session_state.word_doc:
        with st.spinner("Converting PDF to editable document..."):
            try:
                word_doc = pdf_to_docx(pdf_bytes)
                st.session_state.word_doc = word_doc
                st.session_state.show_editor = True
                st.success("✅ Document ready for editing!")
            except Exception as e:
                st.error(f"Conversion failed: {str(e)}")
                st.info("Creating a new blank document instead...")
                word_doc = Document()
                word_doc.add_paragraph("INFORMATION INVOICE")
                word_doc.add_paragraph("")
                st.session_state.word_doc = word_doc
                st.session_state.show_editor = True
    
    # ===== WORD-LIKE EDITOR =====
    if st.session_state.word_doc and st.session_state.show_editor:
        # Toolbar
        st.markdown("""
        <div class="word-toolbar">
            <button>📄 File</button>
            <button>✏️ Edit</button>
            <button>📐 Insert</button>
            <button>🔤 Format</button>
            <span style="margin-left: auto; color: #666; font-size: 13px;">
                💡 Click the text area below to edit
            </span>
        </div>
        """, unsafe_allow_html=True)
        
        # Document content
        st.markdown('<div class="word-editor">', unsafe_allow_html=True)
        
        # Extract content from Word document
        content_lines = []
        for para in st.session_state.word_doc.paragraphs:
            if para.text.strip():
                content_lines.append(para.text)
        
        # Display as textarea for editing
        if 'doc_content' not in st.session_state:
            st.session_state.doc_content = "\n".join(content_lines)
        
        doc_content = st.text_area(
            "",
            value=st.session_state.doc_content,
            height=500,
            key="doc_editor",
            label_visibility="collapsed",
            help="Edit the document content here. Like Microsoft Word!"
        )
        
        # Update content on change
        if doc_content != st.session_state.doc_content:
            st.session_state.doc_content = doc_content
            # Update Word document
            st.session_state.word_doc = Document()
            for line in doc_content.split("\n"):
                if line.strip():
                    p = st.session_state.word_doc.add_paragraph(line.strip())
                    # Check if it's a header
                    if "INVOICE" in line.upper():
                        p.runs[0].bold = True
                        p.runs[0].font.size = Pt(14)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Formatting helpers
        st.markdown("---")
        st.markdown("**Formatting Helpers** (insert these markers for formatting):")
        
        col_fmt1, col_fmt2, col_fmt3, col_fmt4 = st.columns(4)
        with col_fmt1:
            if st.button("📌 **Bold**", use_container_width=True):
                st.session_state.doc_content += " **Bold text** "
                st.rerun()
        with col_fmt2:
            if st.button("📌 *Italic*", use_container_width=True):
                st.session_state.doc_content += " *Italic text* "
                st.rerun()
        with col_fmt3:
            if st.button("📌 Heading", use_container_width=True):
                st.session_state.doc_content += "\n\n=== HEADING ===\n"
                st.rerun()
        with col_fmt4:
            if st.button("📌 New Line", use_container_width=True):
                st.session_state.doc_content += "\n\n"
                st.rerun()
        
        # ===== ACTIONS =====
        st.markdown("---")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            if st.button("📊 Generate Invoice", type="primary", use_container_width=True):
                # Parse document content to extract data
                lines = st.session_state.doc_content.split("\n")
                data = {
                    'guest_name': '',
                    'room_no': '',
                    'num_guests': 1,
                    'arrival': datetime.now().strftime('%d/%m/%y'),
                    'departure': datetime.now().strftime('%d/%m/%y'),
                    'room_type': '',
                    'cashier': '9813FOHALONG',
                    'items': []
                }
                
                # Simple parsing to extract data
                for line in lines:
                    if "Guest:" in line or "Name:" in line:
                        data['guest_name'] = line.split(":")[-1].strip()
                    elif "Room:" in line or "Room No:" in line:
                        data['room_no'] = line.split(":")[-1].strip()
                    elif "Arrival:" in line:
                        data['arrival'] = line.split(":")[-1].strip()
                    elif "Departure:" in line:
                        data['departure'] = line.split(":")[-1].strip()
                    elif "Room Type:" in line:
                        data['room_type'] = line.split(":")[-1].strip()
                
                if data['guest_name']:
                    invoice_pdf = generate_invoice_pdf(data)
                    st.success("✅ Invoice generated!")
                    
                    st.markdown("### 📄 Preview")
                    display_pdf_preview(invoice_pdf, height=500)
                    
                    st.download_button(
                        label="📥 Download Invoice PDF",
                        data=invoice_pdf,
                        file_name=f"Invoice_{data['guest_name'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                else:
                    st.warning("Please enter guest name in the document")
        
        with col2:
            if st.button("📥 Download DOCX", use_container_width=True):
                # Save Word document to bytes
                from io import BytesIO
                docx_bytes = BytesIO()
                st.session_state.word_doc.save(docx_bytes)
                docx_bytes.seek(0)
                
                st.download_button(
                    label="Download Word Document",
                    data=docx_bytes.getvalue(),
                    file_name=f"invoice_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        with col3:
            if st.button("🔄 Reset", use_container_width=True):
                st.session_state.word_doc = None
                st.session_state.doc_content = ""
                st.session_state.show_editor = False
                st.rerun()
        
        with col4:
            st.caption("💡 Edit like Word")